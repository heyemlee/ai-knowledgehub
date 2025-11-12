import logging
from typing import List, Dict, Tuple
from io import BytesIO
import re
import PyPDF2
from docx import Document
import openpyxl
from app.core.constants import DocumentParserConfig

logger = logging.getLogger(__name__)


class DocumentParser:
    @staticmethod
    def parse_pdf(file_content: bytes) -> List[str]:
        """
        解析 PDF 文件
        
        Args:
            file_content: PDF 文件内容（字节）
            
        Returns:
            文本块列表
        """
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            texts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    texts.append(text)
            
            return texts
        except Exception as e:
            logger.error(f"PDF 解析失败: {e}")
            raise
    
    @staticmethod
    def parse_docx(file_content: bytes) -> List[str]:
        """
        解析 Word 文档（支持中文）
        
        Args:
            file_content: DOCX 文件内容（字节）
            
        Returns:
            文本块列表
        """
        try:
            doc_file = BytesIO(file_content)
            doc = Document(doc_file)
            
            texts = []
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    texts.append(para_text)
            
            for table in doc.tables:
                table_texts = []
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        table_texts.append(" | ".join(row_texts))
                
                if table_texts:
                    texts.append("表格内容:\n" + "\n".join(table_texts))
            
            if not texts:
                logger.warning("Word 文档解析后为空，可能文档格式特殊或内容为空")
                texts.append("")
            
            logger.info(f"成功解析 Word 文档，提取了 {len(texts)} 个文本片段")
            return texts
        except Exception as e:
            logger.error(f"DOCX 解析失败: {e}", exc_info=True)
            raise
    
    @staticmethod
    def parse_text(file_content: bytes) -> List[str]:
        """
        解析纯文本文件
        
        Args:
            file_content: TXT 文件内容（字节）
            
        Returns:
            文本块列表
        """
        try:
            text = file_content.decode('utf-8')
            if not text.strip():
                logger.warning("TXT 文件内容为空")
                return [""]
            
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            return lines if lines else [text]
        except UnicodeDecodeError:
            try:
                text = file_content.decode('gbk')
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                return lines if lines else [text]
            except Exception as e:
                logger.error(f"TXT 文件解析失败: {e}")
                raise ValueError(f"无法解析文本文件，编码不支持: {e}")
        except Exception as e:
            logger.error(f"TXT 解析失败: {e}")
            raise
    
    @staticmethod
    def parse_markdown(file_content: bytes) -> List[Dict[str, any]]:
        """
        解析 Markdown 文件，保留结构信息
        
        Args:
            file_content: MD 文件内容（字节）
            
        Returns:
            结构化文本块列表，每个块包含：
            - text: 文本内容
            - heading: 所属标题
            - level: 标题层级 (1-6)
            - section_path: 完整的章节路径
        """
        try:
            # 解码文件内容
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('gbk')
            
            if not text.strip():
                logger.warning("Markdown 文件内容为空")
                return [{"text": "", "heading": "", "level": 0, "section_path": ""}]
            
            # 按行分割
            lines = text.split('\n')
            
            # 解析结构
            sections = []
            current_section = {
                "heading": "",
                "level": 0,
                "content": [],
                "parent_path": []
            }
            heading_stack = []  # 用于跟踪标题层级路径
            
            heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
            
            for line in lines:
                match = heading_pattern.match(line)
                
                if match:
                    # 保存当前段落
                    if current_section["content"]:
                        section_text = '\n'.join(current_section["content"]).strip()
                        if section_text:
                            sections.append({
                                "text": section_text,
                                "heading": current_section["heading"],
                                "level": current_section["level"],
                                "section_path": ' > '.join(current_section["parent_path"])
                            })
                    
                    # 开始新段落
                    level = len(match.group(1))
                    heading = match.group(2).strip()
                    
                    # 更新标题栈
                    while heading_stack and heading_stack[-1]["level"] >= level:
                        heading_stack.pop()
                    
                    heading_stack.append({"level": level, "heading": heading})
                    parent_path = [h["heading"] for h in heading_stack]
                    
                    current_section = {
                        "heading": heading,
                        "level": level,
                        "content": [],
                        "parent_path": parent_path
                    }
                else:
                    # 普通内容行
                    if line.strip():
                        current_section["content"].append(line)
            
            # 保存最后一个段落
            if current_section["content"]:
                section_text = '\n'.join(current_section["content"]).strip()
                if section_text:
                    sections.append({
                        "text": section_text,
                        "heading": current_section["heading"],
                        "level": current_section["level"],
                        "section_path": ' > '.join(current_section["parent_path"])
                    })
            
            # 如果没有解析到任何段落，返回整个文本
            if not sections:
                sections = [{
                    "text": text.strip(),
                    "heading": "",
                    "level": 0,
                    "section_path": ""
                }]
            
            logger.info(f"Markdown 解析完成，识别到 {len(sections)} 个结构化段落")
            return sections
            
        except Exception as e:
            logger.error(f"Markdown 解析失败: {e}")
            raise
    
    @staticmethod
    def parse_excel(file_content: bytes) -> List[str]:
        """
        解析 Excel 文件
        
        Args:
            file_content: Excel 文件内容（字节）
            
        Returns:
            文本块列表
        """
        try:
            excel_file = BytesIO(file_content)
            workbook = openpyxl.load_workbook(excel_file)
            
            texts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                    if row_text.strip():
                        sheet_text.append(row_text)
                if sheet_text:
                    texts.append(f"工作表: {sheet_name}\n" + "\n".join(sheet_text))
            
            return texts
        except Exception as e:
            logger.error(f"Excel 解析失败: {e}")
            raise
    
    @staticmethod
    def chunk_text(
        texts: List[str], 
        chunk_size: int = DocumentParserConfig.DEFAULT_CHUNK_SIZE, 
        overlap: int = DocumentParserConfig.DEFAULT_OVERLAP
    ) -> List[str]:
        """
        将文本切分成块（支持中文，优化结构化信息保留）
        
        Args:
            texts: 文本列表
            chunk_size: 每块大小（字符数，中文按字符计算）
            overlap: 重叠字符数
            
        Returns:
            文本块列表
        """
        if not texts:
            return []
        
        chunks = []
        full_text = "\n\n".join(texts)
        
        start = 0
        text_length = len(full_text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            chunk = full_text[start:end]
            
            if end < text_length:
                best_split = len(chunk)
                
                last_newline = chunk.rfind('\n')
                if last_newline > len(chunk) * DocumentParserConfig.NEWLINE_THRESHOLD:
                    best_split = last_newline + 1
                
                if best_split == len(chunk):
                    for punct in DocumentParserConfig.STRUCTURED_SEPARATORS:
                        punct_pos = chunk.rfind(punct)
                        if punct_pos > len(chunk) * DocumentParserConfig.PUNCTUATION_THRESHOLD:
                            best_split = punct_pos + 1
                            break
                
                if best_split == len(chunk):
                    for punct in DocumentParserConfig.SENTENCE_SEPARATORS:
                        punct_pos = chunk.rfind(punct)
                        if punct_pos > len(chunk) * DocumentParserConfig.SENTENCE_PUNCTUATION_THRESHOLD:
                            best_split = punct_pos + 1
                            break
                
                if best_split < len(chunk) and best_split > len(chunk) * DocumentParserConfig.MIN_SPLIT_RATIO:
                    chunk = chunk[:best_split]
                    end = start + len(chunk)
            
            if chunk.strip():
                chunks.append(chunk)
            
            if end >= text_length:
                break
            
            next_start = end - overlap
            if next_start <= start:
                next_start = start + 1
            
            start = next_start
        
        logger.info(f"文本切分完成: 原始文本长度 {text_length} 字符，切分成 {len(chunks)} 个块")
        return chunks
    
    @staticmethod
    def chunk_markdown(
        sections: List[Dict[str, any]], 
        chunk_size: int = DocumentParserConfig.DEFAULT_CHUNK_SIZE,
        overlap: int = DocumentParserConfig.DEFAULT_OVERLAP
    ) -> Tuple[List[str], List[Dict[str, any]]]:
        """
        智能切分 Markdown 文档，保留语义边界和结构信息
        
        Args:
            sections: parse_markdown 返回的结构化段落列表
            chunk_size: 每块大小（字符数）
            overlap: 重叠字符数
            
        Returns:
            (文本块列表, 元数据列表) 元组，每个元数据包含 heading、level、section_path
        """
        if not sections:
            return [], []
        
        chunks = []
        metadata_list = []
        
        for section in sections:
            text = section["text"]
            heading = section["heading"]
            level = section["level"]
            section_path = section["section_path"]
            
            # 如果段落本身不超过限制，直接作为一个块
            if len(text) <= chunk_size:
                chunks.append(text)
                metadata_list.append({
                    "heading": heading,
                    "level": level,
                    "section_path": section_path
                })
            else:
                # 段落过长，需要切分，但保留段落上下文信息
                start = 0
                text_length = len(text)
                chunk_index = 0
                
                while start < text_length:
                    end = min(start + chunk_size, text_length)
                    chunk = text[start:end]
                    
                    # 寻找最佳切分点
                    if end < text_length:
                        best_split = len(chunk)
                        
                        # 优先在段落边界切分
                        last_newline = chunk.rfind('\n\n')
                        if last_newline > len(chunk) * 0.6:
                            best_split = last_newline + 2
                        
                        # 其次在句子边界切分
                        if best_split == len(chunk):
                            for punct in ['。', '！', '？', '.', '!', '?']:
                                punct_pos = chunk.rfind(punct)
                                if punct_pos > len(chunk) * 0.7:
                                    best_split = punct_pos + 1
                                    break
                        
                        # 最后在单词边界切分
                        if best_split == len(chunk):
                            last_space = chunk.rfind(' ')
                            if last_space > len(chunk) * 0.8:
                                best_split = last_space + 1
                        
                        if best_split < len(chunk) and best_split > len(chunk) * 0.5:
                            chunk = chunk[:best_split]
                            end = start + len(chunk)
                    
                    if chunk.strip():
                        chunks.append(chunk)
                        # 为长段落的每个块添加相同的结构信息
                        metadata_list.append({
                            "heading": heading,
                            "level": level,
                            "section_path": section_path,
                            "section_chunk_index": chunk_index  # 段落内的块序号
                        })
                        chunk_index += 1
                    
                    if end >= text_length:
                        break
                    
                    # 计算下一个起点（带重叠）
                    next_start = end - overlap
                    if next_start <= start:
                        next_start = start + 1
                    
                    start = next_start
        
        logger.info(f"Markdown 智能切块完成: {len(sections)} 个段落 -> {len(chunks)} 个块")
        return chunks, metadata_list

